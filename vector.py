import json
import os

import pandas as pd
from docx import Document as DocxDocument
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_ollama import OllamaEmbeddings

DATA_DIR = "data"
DB_LOCATION = "./chrome_langchain_db"
COLLECTION_NAME = "restaurant_reviews"

_retriever = None


def load_restaurant_text(path: str) -> str:
    with open(path, "rb") as restaurant_file:
        if restaurant_file.read(2) == b"PK":
            docx_file = DocxDocument(path)
            return "\n".join(
                paragraph.text.strip()
                for paragraph in docx_file.paragraphs
                if paragraph.text.strip()
            )

    with open(path, encoding="utf-8") as restaurant_file:
        return restaurant_file.read().strip()


def load_documents() -> tuple[list[Document], list[str]]:
    documents: list[Document] = []
    ids: list[str] = []

    reviews = pd.read_csv(os.path.join(DATA_DIR, "reviews.csv"))
    for i, row in reviews.iterrows():
        documents.append(
            Document(
                page_content=f"{row['Title']}. {row['Review']}",
                metadata={
                    "source": "reviews",
                    "rating": row["Rating"],
                    "date": row["Date"],
                },
            )
        )
        ids.append(str(i))

    restaurant_text = load_restaurant_text(os.path.join(DATA_DIR, "restaurant.docx"))
    documents.append(
        Document(
            page_content=restaurant_text,
            metadata={"source": "restaurant"},
        )
    )
    ids.append("restaurant")

    menu_path = os.path.join(DATA_DIR, "menu.json")
    with open(menu_path, encoding="utf-8") as menu_file:
        menu_items = json.load(menu_file)

    for i, item in enumerate(menu_items):
        ingredients = ", ".join(item["ingredients"])
        documents.append(
            Document(
                page_content=(
                    f"{item['name']} ({item['size']}) - {item['price']} EUR. "
                    f"Ingredients: {ingredients}"
                ),
                metadata={
                    "source": "menu",
                    "name": item["name"],
                    "price": item["price"],
                },
            )
        )
        ids.append(f"menu-{i}")

    return documents, ids


def sync_vector_store() -> Chroma:
    embeddings = OllamaEmbeddings(model="mxbai-embed-large")
    vector_store = Chroma(
        collection_name=COLLECTION_NAME,
        persist_directory=DB_LOCATION,
        embedding_function=embeddings,
    )

    documents, ids = load_documents()
    existing_ids = set(vector_store.get()["ids"])
    new_documents = []
    new_ids = []

    for document, doc_id in zip(documents, ids):
        if doc_id not in existing_ids:
            new_documents.append(document)
            new_ids.append(doc_id)

    if new_documents:
        vector_store.add_documents(documents=new_documents, ids=new_ids)
        print(f"Added {len(new_documents)} new documents to the vector store.")

    return vector_store


def get_retriever():
    global _retriever
    if _retriever is None:
        vector_store = sync_vector_store()
        _retriever = vector_store.as_retriever(search_kwargs={"k": 5})
    return _retriever
